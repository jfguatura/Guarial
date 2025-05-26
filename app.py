import streamlit as st
from docx import Document
from io import BytesIO
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from datetime import datetime
import re

st.set_page_config(page_title="Formatar Arial 16 - Guatura", page_icon="📝")

st.title("📝 Formatar Arial 16 - Guatura")
st.subheader("Ferramenta de formatação rápida para textos")

st.markdown(
    """
    👉 **Passos:**  
    1️⃣ Cole seu texto bruto abaixo.  
    2️⃣ Clique em **'Aplicar Formatação'** para visualizar.  
    3️⃣ Ou clique direto em **'Baixar como .DOCX'** se quiser gerar sem visualizar. 

    _Atenção: O nome do arquivo será salvo com a primeira linha do documento, seguido_
    _da data do dia truncada._
    _Exemplo: "Título - 08112019.docx"_
    """
)

texto = st.text_area("✍️ Cole aqui seu texto bruto", height=300)


def add_page_number(paragraph):
    run = paragraph.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')

    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')

    r_element = run._r
    r_element.append(fldChar1)
    r_element.append(instrText)
    r_element.append(fldChar2)

    font = run.font
    font.size = Pt(10)
    font.name = 'Arial'
    font.color.rgb = RGBColor(150, 150, 150)


def gerar_documento(paragrafos):
    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.18)
        section.right_margin = Inches(1.18)

    for idx, p in enumerate(paragrafos):
        if p.strip():
            par = doc.add_paragraph()
            run = par.add_run(p)
            run.font.name = 'Arial'
            run.font.size = Pt(16)
            if idx == 0:
                run.bold = True
            par.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            par.paragraph_format.space_before = Pt(12)
            par.paragraph_format.space_after = Pt(12)
            par.paragraph_format.line_spacing = 1.0

    section = doc.sections[0]
    footer = section.footer

    for p in footer.paragraphs:
        p.clear()

    footer_paragraph = footer.add_paragraph()

    left_run = footer_paragraph.add_run("GPPR/JFG")
    left_run.font.size = Pt(10)
    left_run.font.name = 'Arial'
    left_run.font.color.rgb = RGBColor(150, 150, 150)

    tab_stops = footer_paragraph.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Inches(6.0))

    footer_paragraph.add_run('\t')
    add_page_number(footer_paragraph)

    footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

    return doc


# 🔥 Processa o texto sempre que houver algo digitado
if texto.strip():
    texto_formatado = texto.replace("\n\n", "\n")
    paragrafos = texto_formatado.split('\n')

    titulo = paragrafos[0].strip() if paragrafos and paragrafos[0].strip() else "Texto_Formatado"
    titulo_seguro = re.sub(r'[\/:*?"<>|()]+', '-', titulo)
    titulo_seguro = re.sub(r'\s+', ' ', titulo_seguro).strip()

    data_atual = datetime.now().strftime("%d%m%y")
    nome_arquivo = f"{titulo_seguro} - {data_atual}.docx"

    doc = gerar_documento(paragrafos)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    col1, col2 = st.columns([1, 1])

    with col1:
        if st.button("✅ Aplicar Formatação"):
            st.markdown("### 🎨 Texto Formatado (Visualização):")
            for idx, p in enumerate(paragrafos):
                if p.strip():
                    negrito = "font-weight: bold;" if idx == 0 else ""
                    st.markdown(
                        f"""<div style='font-family: Arial; font-size: 16px; text-align: justify; margin: 12px 0; {negrito}'>{p}</div>""",
                        unsafe_allow_html=True
                    )

    with col2:
        st.download_button(
            label="📥 Baixar como .DOCX",
            data=buffer,
            file_name=nome_arquivo,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

else:
    st.info("ℹ️ Insira um texto acima para visualizar ou gerar o documento.")
    

st.markdown("---")
st.markdown("💡 **Observação:** O rodapé inclui texto à esquerda e número de página automático (campo Word). Se o número não aparecer, pressione 'F9' no Word para atualizar os campos.")
